import pdfplumber
from docx import Document
from typing import Dict, Any, Optional, Tuple
import re
import json
from pathlib import Path
from ..models.schemas import ResumeData, PersonalInfo, ExperienceEntry, EducationEntry, BulletPoint

# Import LLM function
try:
    from ..config import generate_chat_completion
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    print("Warning: LLM not available for resume parsing")

class ResumeParser:
    """Parse resumes from PDF or DOCX format into structured data."""
    
    def __init__(self):
        self.section_keywords = {
            'experience': ['experience', 'work experience', 'employment', 'employment history', 'professional experience'],
            'education': ['education', 'academic', 'qualifications', 'degrees'],
            'skills': ['skills', 'technical skills', 'competencies', 'expertise', 'proficiencies', 'core competencies'],
            'summary': ['summary', 'professional summary', 'profile', 'objective', 'about'],
            'certifications': ['certifications', 'certificates', 'licenses', 'credentials', 'professional certifications']
        }
    
    def parse(self, file_path: str) -> Tuple[ResumeData, Dict[str, Any]]:
        """
        Parse resume file and return structured data + formatting map.
        
        Args:
            file_path: Path to resume file (PDF or DOCX)
            
        Returns:
            Tuple of (ResumeData, formatting_map)
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        try:
            if extension == '.pdf':
                result = self._parse_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                result = self._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            
            # Validate return value
            if not isinstance(result, tuple):
                raise ValueError(f"Parser returned {type(result)}, expected tuple")
            if len(result) != 2:
                raise ValueError(f"Parser returned {len(result)} values, expected 2")
            
            return result
        except ValueError:
            # Re-raise ValueError as-is
            raise
        except Exception as e:
            raise ValueError(f"Error parsing resume file: {str(e)}")
    
    def _parse_pdf(self, file_path: str) -> Tuple[ResumeData, Dict[str, Any]]:
        """Parse PDF resume."""
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                    else:
                        print(f"Warning: Page {page_num + 1} had no extractable text")
            
            full_text = "\n".join(text_content)
            
            if not full_text or not full_text.strip():
                raise ValueError("PDF file appears to be empty or contains no extractable text. The PDF might be image-based or encrypted.")
            
            # Debug: print raw text
            print(f"\n=== Raw PDF Text (first 1500 chars) ===")
            print(full_text[:1500])
            print("=" * 40)
            
            # Try LLM parsing first for best accuracy
            if LLM_AVAILABLE:
                try:
                    resume_data = self._parse_with_llm(full_text)
                    if resume_data:
                        print("✓ LLM parsing successful")
                        return resume_data, {}
                except Exception as e:
                    print(f"LLM parsing failed, falling back to regex: {e}")
            
            # Fallback to regex parsing
            full_text = self._join_wrapped_lines(full_text)
            resume_data = self._parse_text(full_text, {})
            return resume_data, {}
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Error reading PDF file: {str(e)}")
    
    def _parse_with_llm(self, text: str) -> Optional[ResumeData]:
        """Use LLM to parse resume text into structured data."""
        prompt = f"""Extract information from this resume into JSON. 

=== CRITICAL RULES ===
1. DO NOT make up or invent ANY information
2. Extract EXACTLY what's written in the resume

RESUME TEXT:
{text}

=== OUTPUT FORMAT ===
Return JSON with this EXACT structure:
{{
    "personal_info": {{
        "name": "full name",
        "email": "email@example.com", 
        "phone": "phone number",
        "location": "city, state",
        "linkedin": "linkedin.com/in/profile"
    }},
    "summary": "professional summary text",
    "experience": [
        {{
            "company": "company name",
            "title": "job title",
            "location": "city, state",
            "start_date": "start date",
            "end_date": "end date",
            "bullets": ["bullet point 1", "bullet point 2"]
        }}
    ],
    "education": [
        {{
            "institution": "school name",
            "degree": "degree",
            "field": "field of study",
            "location": "city, state",
            "graduation_date": "date"
        }}
    ],
    "skills": ["skill1", "skill2", "skill3"]
}}

=== SKILLS EXTRACTION ===
Put ALL skills, certifications, tools, technologies, and competencies into the "skills" array.
- Include everything from "Skills", "Technical Skills", "Certifications" sections
- Split comma-separated items into individual entries
- Extract each item exactly as written

Return ONLY valid JSON, no markdown or explanations."""

        try:
            response = generate_chat_completion(
                model="grok-4-fast",
                messages=[
                    {"role": "system", "content": "You are an expert resume parser. Extract all information accurately and return valid JSON only."},
                    {"role": "user", "content": prompt}
                ]
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # Clean JSON
            if result_text.startswith("```"):
                result_text = re.sub(r'^```(?:json)?\s*', '', result_text)
                result_text = re.sub(r'\s*```$', '', result_text)
            
            data = json.loads(result_text)
            
            # Build ResumeData from parsed JSON
            personal_info = PersonalInfo(
                name=data.get('personal_info', {}).get('name'),
                email=data.get('personal_info', {}).get('email'),
                phone=data.get('personal_info', {}).get('phone'),
                location=data.get('personal_info', {}).get('location'),
                linkedin=data.get('personal_info', {}).get('linkedin')
            )
            
            # Parse experience
            experience = []
            for exp in data.get('experience', []):
                bullets = []
                for b in exp.get('bullets', []):
                    if isinstance(b, str) and b.strip():
                        bullets.append(BulletPoint(text=b.strip()))
                
                if bullets:  # Only add if there are bullets
                    experience.append(ExperienceEntry(
                        company=exp.get('company', 'Unknown'),
                        title=exp.get('title', 'Unknown'),
                        location=exp.get('location'),
                        start_date=exp.get('start_date'),
                        end_date=exp.get('end_date'),
                        bullets=bullets
                    ))
            
            # Parse education
            education = []
            for edu in data.get('education', []):
                education.append(EducationEntry(
                    institution=edu.get('institution', 'Unknown'),
                    degree=edu.get('degree', 'Unknown'),
                    field=edu.get('field'),
                    location=edu.get('location'),
                    graduation_date=edu.get('graduation_date')
                ))
            
            # Get skills from LLM
            raw_skills = data.get('skills', [])
            
            print(f"\n" + "="*60)
            print(f"=== RAW LLM OUTPUT ===")
            print(f"Raw skills from LLM: {raw_skills}")
            print("="*60)
            
            # Process skills - flatten comma-separated items
            skills = []
            for item in raw_skills:
                if not isinstance(item, str) or not item.strip():
                    continue
                item_clean = item.strip()
                # Remove any prefix labels
                item_clean = re.sub(r'^(Skills?|Certifications?|Technical Skills?)\s*:\s*', '', item_clean, flags=re.IGNORECASE)
                # Split by comma
                if ',' in item_clean:
                    sub_items = [s.strip() for s in item_clean.split(',') if s.strip()]
                else:
                    sub_items = [item_clean] if item_clean else []
                for sub_item in sub_items:
                    if sub_item and sub_item not in skills:
                        skills.append(sub_item)
            
            print(f"\n=== FINAL RESULT ===")
            print(f"Skills ({len(skills)}): {skills}")
            
            return ResumeData(
                personal_info=personal_info,
                summary=data.get('summary'),
                experience=experience,
                education=education,
                skills=skills,
                formatting_map=None
            )
            
        except Exception as e:
            print(f"LLM parsing error: {e}")
            return None
    
    def _join_wrapped_lines(self, text: str) -> str:
        """Join lines that are continuations of previous lines (wrapped text in PDF)."""
        lines = text.split('\n')
        joined_lines = []
        current_line = ""
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current_line:
                    joined_lines.append(current_line)
                    current_line = ""
                continue
            
            # Check if this line starts a new logical section/bullet/entry
            is_new_item = (
                stripped.startswith('•') or
                stripped.startswith('- ') or
                stripped.startswith('* ') or
                re.match(r'^\d+\.\s', stripped) or
                re.match(r'^[A-Z][A-Z\s&]+$', stripped) or
                re.match(r'^[A-Z][A-Za-z\s]+,\s*[A-Z]', stripped) or
                re.match(r'.*\|\s*\w+\s+\d{4}', stripped) or
                re.match(r'^(WORK\s+)?EXPERIENCE|EDUCATION|SKILLS|SUMMARY|CERTIFICATIONS|TECHNICAL', stripped, re.IGNORECASE) or
                re.match(r'^(Senior\s+|Lead\s+|Staff\s+|Principal\s+)?(Product|Project|Program|Engineering|Software|Data|Technical)\s+(Manager|Director|Engineer|Analyst|Lead)', stripped, re.IGNORECASE)
            )
            
            if is_new_item:
                if current_line:
                    joined_lines.append(current_line)
                current_line = stripped
            else:
                if current_line:
                    if current_line.startswith('•') or current_line.startswith('- '):
                        current_line += " " + stripped
                    elif not current_line.endswith(('.', ':', '!', '?')):
                        current_line += " " + stripped
                    else:
                        joined_lines.append(current_line)
                        current_line = stripped
                else:
                    current_line = stripped
        
        if current_line:
            joined_lines.append(current_line)
        
        return '\n'.join(joined_lines)
    
    def _parse_docx(self, file_path: str) -> Tuple[ResumeData, Dict[str, Any]]:
        """Parse DOCX resume with formatting preservation."""
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}. The file might be corrupted or in an unsupported format.")
        
        # Extract all text
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Try LLM parsing first
        if LLM_AVAILABLE:
            try:
                resume_data = self._parse_with_llm(full_text)
                if resume_data:
                    print("✓ LLM parsing successful for DOCX")
                    return resume_data, {}
            except Exception as e:
                print(f"LLM parsing failed for DOCX, falling back to regex: {e}")
        
        # Fallback to regex parsing
        formatting_map = {}
        sections = {}
        current_section = None
        current_text = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            section_type = self._identify_section(text)
            if section_type:
                if current_section and current_text:
                    sections[current_section] = "\n".join(current_text)
                current_section = section_type
                current_text = []
                continue
            
            formatting_map[f"para_{i}"] = {
                'text': text,
                'style': para.style.name if para.style else None,
                'runs': [{
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'font_size': run.font.size.pt if run.font.size else None,
                    'font_name': run.font.name
                } for run in para.runs]
            }
            
            if current_section:
                current_text.append(text)
        
        if current_section and current_text:
            sections[current_section] = "\n".join(current_text)
        
        resume_data = self._parse_text(full_text, sections)
        return resume_data, formatting_map
    
    def _parse_text(self, text: str, sections: Dict[str, str]) -> ResumeData:
        """Parse text content into structured resume data (regex fallback)."""
        if not text or not text.strip():
            raise ValueError("Resume text is empty.")
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            raise ValueError("No content found in resume.")
        
        personal_info = self._extract_personal_info(lines[:15])
        summary = self._extract_summary(text, sections)
        experience = self._extract_experience(text, sections)
        education = self._extract_education(text, sections)
        skills = self._extract_skills(text, sections)
        # Also extract certifications and add them to skills
        certifications = self._extract_certifications(text, sections)
        all_skills = skills + [c for c in certifications if c not in skills]
        
        if not personal_info.name and not personal_info.email and len(experience) == 0 and len(education) == 0:
            raise ValueError("Could not extract meaningful content from resume.")
        
        return ResumeData(
            personal_info=personal_info,
            summary=summary,
            experience=experience,
            education=education,
            skills=all_skills,
            formatting_map=None
        )
    
    def _identify_section(self, text: str) -> Optional[str]:
        """Identify if a line is a section header."""
        text_lower = text.lower()
        for section_type, keywords in self.section_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                return section_type
        return None
    
    def _extract_personal_info(self, lines: list[str]) -> PersonalInfo:
        """Extract personal information from header lines."""
        info = PersonalInfo()
        full_text = " ".join(lines)
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, full_text)
        if emails:
            info.email = emails[0]
        
        # Extract phone
        phone_patterns = [
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        ]
        for pattern in phone_patterns:
            phones = re.findall(pattern, full_text)
            if phones:
                info.phone = phones[0]
                break
        
        # Extract LinkedIn
        linkedin_patterns = [
            r'linkedin\.com/in/[\w-]+',
            r'linkedin\.com/[\w-]+',
            r'www\.linkedin\.com/in/[\w-]+'
        ]
        for pattern in linkedin_patterns:
            linkedin = re.findall(pattern, full_text, re.IGNORECASE)
            if linkedin:
                info.linkedin = linkedin[0]
                break
        
        # Extract location (City, State pattern)
        location_patterns = [
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2}(?:\s+\d{5})?)',
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+)'
        ]
        for pattern in location_patterns:
            locations = re.findall(pattern, full_text)
            if locations:
                # Filter out false positives
                for loc in locations:
                    if not any(kw in loc.lower() for kw in ['university', 'college', 'inc', 'llc', 'corp']):
                        info.location = loc
                        break
                break
        
        # Extract name (usually first line, capitalized)
        if lines:
            potential_name = lines[0].strip()
            # Name should be 2-4 words, all capitalized or title case
            if len(potential_name.split()) <= 4 and '@' not in potential_name and 'linkedin' not in potential_name.lower():
                if re.match(r'^[A-Z][a-zA-Z]+(\s+[A-Z][a-zA-Z]+)+$', potential_name) or potential_name.isupper():
                    info.name = potential_name
        
        return info
    
    def _extract_summary(self, text: str, sections: Dict[str, str]) -> Optional[str]:
        """Extract professional summary."""
        if 'summary' in sections:
            return sections['summary']
        
        summary_patterns = [
            r'(?i)(?:summary|professional summary|profile|objective)[:\s]*(.+?)(?=\n\s*(?:WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT|EDUCATION|SKILLS|TECHNICAL\s+SKILLS|CERTIFICATIONS)\b)',
            r'(?i)(?:summary|professional summary|profile|objective)[:\s]*(.+?)(?=\n\n)',
        ]
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                summary_text = match.group(1).strip()
                if len(summary_text) < 2000:
                    return summary_text
        
        return None
    
    def _extract_experience(self, text: str, sections: Dict[str, str]) -> list[ExperienceEntry]:
        """Extract work experience entries."""
        experience_entries = []
        
        if 'experience' in sections:
            exp_text = sections['experience']
        else:
            exp_patterns = [
                r'(?i)(?:WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT(?:\s+HISTORY)?)[:\s]*\n(.+?)(?=\n\s*(?:TECHNICAL\s+SKILLS|SKILLS|EDUCATION|CERTIFICATIONS)\b)',
                r'(?i)(?:WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT)[:\s]*\n(.+?)(?=\n\n\s*[A-Z])',
            ]
            exp_text = None
            for pattern in exp_patterns:
                exp_match = re.search(pattern, text, re.DOTALL)
                if exp_match:
                    exp_text = exp_match.group(1)
                    break
            
            if not exp_text:
                return []
        
        lines = exp_text.split('\n')
        current_entry = []
        entries = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            is_header = bool(re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b', line, re.IGNORECASE))
            is_header = is_header or bool(re.search(r'\|\s*\w+\s+\d{4}\s*[-–]', line))
            is_header = is_header or bool(re.search(r'\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)', line, re.IGNORECASE))
            
            is_potential_header = not line.startswith('•') and not line.startswith('-') and len(line) > 10
            
            if is_header and current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = [line]
            elif is_potential_header and is_header:
                if current_entry:
                    entries.append('\n'.join(current_entry))
                current_entry = [line]
            else:
                current_entry.append(line)
        
        if current_entry:
            entries.append('\n'.join(current_entry))
        
        for entry_text in entries:
            if not entry_text.strip():
                continue
            
            lines = [l.strip() for l in entry_text.split('\n') if l.strip()]
            if not lines:
                continue
            
            first_line = lines[0]
            company = first_line
            title = first_line
            
            if len(lines) > 1 and not lines[1].startswith('•') and not lines[1].startswith('-'):
                if re.search(r'\|.*\d{4}', first_line):
                    company = re.sub(r'\s*\|.*$', '', first_line).strip()
                    title = lines[1]
                else:
                    title = first_line
                    company = lines[1] if len(lines) > 1 else first_line
            
            date_pattern = r'(\w+\s+\d{4}|\d{4})\s*[-–]\s*(\w+\s+\d{4}|\d{4}|present|current)'
            dates = re.findall(date_pattern, entry_text, re.IGNORECASE)
            
            start_date = dates[0][0] if dates else None
            end_date = dates[0][1] if dates else None
            
            bullets = []
            for line in lines:
                if re.match(r'^[•\-\*]\s*', line):
                    bullet_text = re.sub(r'^[•\-\*]\s*', '', line).strip()
                    if bullet_text and len(bullet_text) > 5:
                        bullets.append(BulletPoint(text=bullet_text))
            
            if bullets:
                experience_entries.append(ExperienceEntry(
                    company=company,
                    title=title,
                    start_date=start_date,
                    end_date=end_date,
                    bullets=bullets
                ))
        
        return experience_entries
    
    def _extract_education(self, text: str, sections: Dict[str, str]) -> list[EducationEntry]:
        """Extract education entries."""
        education_entries = []
        
        if 'education' in sections:
            edu_text = sections['education']
        else:
            edu_patterns = [
                r'(?i)\bEDUCATION\b[:\s]*\n(.+?)(?=\n\s*(?:SKILLS|EXPERIENCE|CERTIFICATIONS|PROJECTS)\b|$)',
                r'(?i)\bEDUCATION\b[:\s]*(.+?)(?=\n\n)',
            ]
            edu_text = None
            for pattern in edu_patterns:
                edu_match = re.search(pattern, text, re.DOTALL)
                if edu_match:
                    edu_text = edu_match.group(1)
                    break
            
            if not edu_text:
                return []
        
        lines = edu_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            line = re.sub(r'^[•\-\*]\s*', '', line).strip()
            if not line or len(line) < 5:
                continue
            
            if line.lower().startswith(('skills', 'technical', 'certification')):
                continue
            
            degree_keywords = ['Master', 'Bachelor', 'PhD', 'Doctor', 'MBA', 'MS', 'BS', 'BA', 'Associate', 'Diploma', 'Certificate', 'Program']
            has_degree = any(kw.lower() in line.lower() for kw in degree_keywords)
            
            if has_degree:
                date_match = re.search(r'(\d{4}|Expected\s+\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})', line, re.IGNORECASE)
                graduation_date = date_match.group(1) if date_match else None
                
                parts = [p.strip() for p in line.split(',')]
                
                degree = parts[0] if parts else line
                institution = ""
                
                for i, part in enumerate(parts):
                    if any(kw in part for kw in ['University', 'College', 'Institute', 'Academy', 'School']):
                        institution = part.strip()
                        break
                
                if not institution and len(parts) > 1:
                    institution = parts[-2] if graduation_date and len(parts) > 2 else parts[-1]
                    institution = re.sub(r'\d{4}|Expected|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec', '', institution, flags=re.IGNORECASE).strip().strip(',')
                
                if institution or degree:
                    education_entries.append(EducationEntry(
                        institution=institution or "Unknown",
                        degree=degree,
                        graduation_date=graduation_date
                    ))
        
        return education_entries
    
    def _extract_skills(self, text: str, sections: Dict[str, str]) -> list[str]:
        """Extract skills list (excluding certifications)."""
        if 'skills' in sections:
            skills_text = sections['skills']
        else:
            skills_patterns = [
                r'(?i)(?:TECHNICAL\s+SKILLS|SKILLS\s*(?:&\s*COMPETENCIES)?)[^:]*[:\s]*\n(.+?)(?=\n\s*(?:EDUCATION|CERTIFICATIONS?)\b)',
                r'(?i)(?:TECHNICAL\s+SKILLS|SKILLS)[^:]*[:\s]*(.+?)(?=\n\s*(?:EDUCATION|CERTIFICATIONS?)\b)',
                r'(?i)(?:TECHNICAL\s+SKILLS|SKILLS)[:\s]*(.+?)(?=\n\n)',
            ]
            skills_text = None
            for pattern in skills_patterns:
                skills_match = re.search(pattern, text, re.DOTALL)
                if skills_match:
                    skills_text = skills_match.group(1)
                    break
            
            if not skills_text:
                return []
        
        skills_text = re.sub(r'(?i)\b(?:EDUCATION|CERTIFICATIONS?)\b.*$', '', skills_text, flags=re.DOTALL)
        
        skills = []
        lines = skills_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if any(kw in line for kw in ['University', 'College', 'Master', 'Bachelor', 'Expected', 'Degree']):
                continue
            
            if re.search(r'(?i)\b(?:certified|certification|certificate|PMP|AWS\s+Certified|Google\s+Certified|Scrum\s+Master|Professional|License)\b', line):
                continue
            
            line = re.sub(r'^[•\-\*]\s*', '', line).strip()
            
            if line.startswith('&') or line.lower() in ['certifications', '& certifications']:
                continue
            
            for item in re.split(r'[,;]', line):
                skill = item.strip()
                skill = re.sub(r'^(?:Certifications?:?\s*)', '', skill, flags=re.IGNORECASE)
                
                if re.search(r'(?i)\b(?:certified|certification|certificate|PMP|license)\b', skill):
                    continue
                
                if skill and len(skill) > 1 and len(skill) < 50:
                    if not any(kw in skill for kw in ['University', 'College', 'Master', 'Bachelor', 'Expected']):
                        skills.append(skill)
        
        return skills
    
    def _extract_certifications(self, text: str, sections: Dict[str, str]) -> list[str]:
        """Extract certifications list."""
        if 'certifications' in sections:
            cert_text = sections['certifications']
        else:
            cert_patterns = [
                r'(?i)(?:CERTIFICATIONS?|CERTIFICATES?|LICENSES?)[:\s]*\n(.+?)(?=\n\s*(?:EDUCATION|SKILLS|EXPERIENCE)\b|$)',
                r'(?i)(?:CERTIFICATIONS?|CERTIFICATES?)[:\s]*(.+?)(?=\n\n|$)',
            ]
            cert_text = None
            for pattern in cert_patterns:
                cert_match = re.search(pattern, text, re.DOTALL)
                if cert_match:
                    cert_text = cert_match.group(1)
                    break
            
            if not cert_text:
                return []
        
        certifications = []
        lines = cert_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            line = re.sub(r'^[•\-\*]\s*', '', line).strip()
            
            if re.match(r'^(?:EDUCATION|SKILLS|EXPERIENCE)', line, re.IGNORECASE):
                break
            
            if line and len(line) > 3 and len(line) < 150:
                certifications.append(line)
        
        return certifications
