from docx import Document
from pathlib import Path
from typing import Dict, List, Optional
from ..models.schemas import ResumeData, ExperienceEntry, BulletPoint

# PDF generation imports
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.colors import HexColor, black, gray
from reportlab.lib import colors

class ResumeBuilder:
    """Rebuild resume DOCX/PDF with accepted changes while preserving formatting."""
    
    def __init__(self):
        self.styles = self._create_styles()
    
    def _create_styles(self):
        """Create custom styles for the PDF resume."""
        styles = getSampleStyleSheet()
        
        # Name style - large, bold, centered
        styles.add(ParagraphStyle(
            name='ResumeName',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=6,
            textColor=HexColor('#1a1a2e'),
            fontName='Helvetica-Bold'
        ))
        
        # Contact info style - centered, smaller
        styles.add(ParagraphStyle(
            name='ContactInfo',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            spaceAfter=12,
            textColor=HexColor('#4a4a4a')
        ))
        
        # Section header style
        styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=14,
            spaceAfter=8,
            textColor=HexColor('#2d3436'),
            fontName='Helvetica-Bold',
            borderPadding=(0, 0, 3, 0),
            borderWidth=0,
            borderColor=HexColor('#667eea')
        ))
        
        # Job title style
        styles.add(ParagraphStyle(
            name='JobTitle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            textColor=HexColor('#2d3436'),
            spaceAfter=2
        ))
        
        # Company/date style
        styles.add(ParagraphStyle(
            name='CompanyDate',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#636e72'),
            spaceAfter=4
        ))
        
        # Bullet point style
        styles.add(ParagraphStyle(
            name='BulletPoint',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=15,
            bulletIndent=5,
            spaceAfter=4,
            alignment=TA_JUSTIFY,
            textColor=HexColor('#2d3436'),
            leading=14
        ))
        
        # Summary style
        styles.add(ParagraphStyle(
            name='Summary',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
            textColor=HexColor('#2d3436'),
            leading=14
        ))
        
        # Skills style
        styles.add(ParagraphStyle(
            name='Skills',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            textColor=HexColor('#2d3436')
        ))
        
        return styles
    
    def rebuild_resume_pdf(
        self,
        resume_data: ResumeData,
        accepted_changes: Dict[str, str],
        output_path: str
    ) -> str:
        """
        Generate a PDF resume with accepted changes applied.
        
        Args:
            resume_data: The parsed resume data
            accepted_changes: Dict mapping section_id to new text
            output_path: Path to save the PDF
            
        Returns:
            Path to generated PDF file
        """
        # Apply changes to resume_data
        modified_data = self._apply_changes_to_data(resume_data, accepted_changes)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=0.6*inch,
            leftMargin=0.6*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        story = []
        
        # Personal Info / Header
        if modified_data.personal_info.name:
            story.append(Paragraph(modified_data.personal_info.name, self.styles['ResumeName']))
        
        # Contact line
        contact_parts = []
        if modified_data.personal_info.email:
            contact_parts.append(modified_data.personal_info.email)
        if modified_data.personal_info.phone:
            contact_parts.append(modified_data.personal_info.phone)
        if modified_data.personal_info.location:
            contact_parts.append(modified_data.personal_info.location)
        if modified_data.personal_info.linkedin:
            contact_parts.append(modified_data.personal_info.linkedin)
        
        if contact_parts:
            contact_text = "  |  ".join(contact_parts)
            story.append(Paragraph(contact_text, self.styles['ContactInfo']))
        
        # Divider line
        story.append(Spacer(1, 5))
        story.append(self._create_divider())
        
        # Summary Section
        if modified_data.summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", self.styles['SectionHeader']))
            # Handle multi-line summary
            summary_lines = modified_data.summary.split('\n')
            for line in summary_lines:
                line = line.strip()
                if line:
                    if line.startswith('•') or line.startswith('-'):
                        story.append(Paragraph(f"• {line.lstrip('•- ')}", self.styles['BulletPoint']))
                    else:
                        story.append(Paragraph(line, self.styles['Summary']))
        
        # Experience Section
        if modified_data.experience:
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", self.styles['SectionHeader']))
            
            for exp_idx, exp in enumerate(modified_data.experience):
                # Job title and dates on same line
                title_text = f"<b>{exp.title}</b>"
                if exp.start_date or exp.end_date:
                    date_range = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
                    title_text = f"<b>{exp.title}</b> <font color='#636e72'>| {date_range}</font>"
                
                story.append(Paragraph(title_text, self.styles['JobTitle']))
                
                # Company
                if exp.company:
                    story.append(Paragraph(exp.company, self.styles['CompanyDate']))
                
                # Bullets
                if exp.bullets:
                    for bullet_idx, bullet in enumerate(exp.bullets):
                        bullet_text = bullet.text if isinstance(bullet, BulletPoint) else str(bullet)
                        if bullet_text:
                            story.append(Paragraph(f"• {bullet_text}", self.styles['BulletPoint']))
                
                story.append(Spacer(1, 8))
        
        # Education Section
        if modified_data.education:
            story.append(Paragraph("EDUCATION", self.styles['SectionHeader']))
            
            for edu in modified_data.education:
                edu_text = f"<b>{edu.degree}</b>"
                if edu.graduation_date:
                    edu_text += f" <font color='#636e72'>| {edu.graduation_date}</font>"
                story.append(Paragraph(edu_text, self.styles['JobTitle']))
                
                if edu.institution:
                    story.append(Paragraph(edu.institution, self.styles['CompanyDate']))
                
                story.append(Spacer(1, 4))
        
        # Skills Section
        if modified_data.skills:
            story.append(Paragraph("SKILLS", self.styles['SectionHeader']))
            
            # Group skills into rows of reasonable length
            skills_text = "  •  ".join(modified_data.skills)
            story.append(Paragraph(skills_text, self.styles['Skills']))
        
        # Build PDF
        doc.build(story)
        return output_path
    
    def _create_divider(self):
        """Create a horizontal divider line."""
        from reportlab.platypus import HRFlowable
        return HRFlowable(
            width="100%",
            thickness=1,
            color=HexColor('#667eea'),
            spaceBefore=2,
            spaceAfter=10
        )
    
    def _apply_changes_to_data(
        self,
        resume_data: ResumeData,
        accepted_changes: Dict[str, str]
    ) -> ResumeData:
        """Apply accepted changes to resume data."""
        import copy
        modified = copy.deepcopy(resume_data)
        
        for section_id, new_text in accepted_changes.items():
            if section_id == 'summary':
                modified.summary = new_text
            elif section_id.startswith('experience_'):
                # Parse: experience_{exp_idx}_bullet_{bullet_idx}
                parts = section_id.split('_')
                if len(parts) >= 4 and parts[2] == 'bullet':
                    try:
                        exp_idx = int(parts[1])
                        bullet_idx = int(parts[3])
                        if exp_idx < len(modified.experience):
                            exp = modified.experience[exp_idx]
                            if bullet_idx < len(exp.bullets):
                                if isinstance(exp.bullets[bullet_idx], BulletPoint):
                                    exp.bullets[bullet_idx].text = new_text
                                else:
                                    exp.bullets[bullet_idx] = BulletPoint(text=new_text)
                    except (ValueError, IndexError) as e:
                        print(f"Error applying change to {section_id}: {e}")
            elif section_id.startswith('skill'):
                # Add new skill if it doesn't exist
                if new_text not in modified.skills:
                    new_skills = [s.strip() for s in new_text.split(',')]
                    for skill in new_skills:
                        if skill and skill not in modified.skills:
                            modified.skills.append(skill)
        
        return modified
    
    def rebuild_resume(
        self,
        original_path: str,
        resume_data: ResumeData,
        accepted_changes: Dict[str, str],
        output_path: str
    ) -> str:
        """
        Rebuild resume DOCX with accepted changes.
        
        Args:
            original_path: Path to original DOCX file
            resume_data: Updated resume data with accepted changes
            accepted_changes: Dict mapping section_id to new text
            output_path: Path to save the new resume
            
        Returns:
            Path to generated resume file
        """
        try:
            # Load original document
            doc = Document(original_path)
            
            # Apply changes to document
            self._apply_changes_to_document(doc, resume_data, accepted_changes)
            
            # Save new document
            doc.save(output_path)
            return output_path
        except Exception as e:
            # If DOCX manipulation fails, create new document
            return self._create_new_document(resume_data, output_path)
    
    def _apply_changes_to_document(
        self,
        doc: Document,
        resume_data: ResumeData,
        accepted_changes: Dict[str, str]
    ):
        """Apply accepted changes to existing document structure."""
        # Map section IDs to document locations
        para_index = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this paragraph matches any changed section
            for section_id, new_text in accepted_changes.items():
                if self._paragraph_matches_section(para, section_id, resume_data):
                    # Replace text while preserving formatting
                    self._replace_paragraph_text(para, new_text)
                    break
            
            para_index += 1
    
    def _paragraph_matches_section(self, para, section_id: str, resume_data: ResumeData) -> bool:
        """Check if paragraph matches a section ID."""
        para_text = para.text.lower()
        
        if section_id == 'summary_section':
            return resume_data.summary and resume_data.summary.lower() in para_text
        elif section_id == 'skills_section':
            return any(skill.lower() in para_text for skill in resume_data.skills)
        elif section_id.startswith('experience_'):
            # Parse section_id: experience_{idx}_bullet_{bullet_idx}
            parts = section_id.split('_')
            if len(parts) >= 2:
                try:
                    exp_idx = int(parts[1])
                    if exp_idx < len(resume_data.experience):
                        exp = resume_data.experience[exp_idx]
                        # Check if paragraph contains company or title
                        return (exp.company.lower() in para_text or 
                                exp.title.lower() in para_text)
                except (ValueError, IndexError):
                    pass
        
        return False
    
    def _replace_paragraph_text(self, para, new_text: str):
        """Replace paragraph text while preserving formatting of first run."""
        if not para.runs:
            para.add_run(new_text)
            return
        
        # Keep formatting from first run
        first_run = para.runs[0]
        para.clear()
        
        new_run = para.add_run(new_text)
        new_run.bold = first_run.bold
        new_run.italic = first_run.italic
        if first_run.font.size:
            new_run.font.size = first_run.font.size
        if first_run.font.name:
            new_run.font.name = first_run.font.name
        if first_run.font.color:
            new_run.font.color = first_run.font.color
    
    def _create_new_document(self, resume_data: ResumeData, output_path: str) -> str:
        """Create a new DOCX document from resume data."""
        doc = Document()
        
        # Add personal info
        if resume_data.personal_info.name:
            heading = doc.add_heading(resume_data.personal_info.name, level=1)
            heading.alignment = 1  # Center alignment
        
        if resume_data.personal_info.email:
            doc.add_paragraph(resume_data.personal_info.email)
        if resume_data.personal_info.phone:
            doc.add_paragraph(resume_data.personal_info.phone)
        
        doc.add_paragraph()  # Blank line
        
        # Add summary
        if resume_data.summary:
            doc.add_heading('Professional Summary', level=2)
            doc.add_paragraph(resume_data.summary)
            doc.add_paragraph()
        
        # Add experience
        if resume_data.experience:
            doc.add_heading('Experience', level=2)
            for exp in resume_data.experience:
                # Company and title
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{exp.title} - {exp.company}")
                title_run.bold = True
                
                if exp.start_date or exp.end_date:
                    date_str = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
                    doc.add_paragraph(date_str, style='Intense Quote')
                
                # Bullets
                for bullet in exp.bullets:
                    bullet_text = bullet.text if isinstance(bullet, BulletPoint) else str(bullet)
                    bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            
            doc.add_paragraph()
        
        # Add education
        if resume_data.education:
            doc.add_heading('Education', level=2)
            for edu in resume_data.education:
                edu_para = doc.add_paragraph()
                edu_para.add_run(f"{edu.degree}").bold = True
                edu_para.add_run(f" - {edu.institution}")
                if edu.graduation_date:
                    doc.add_paragraph(edu.graduation_date, style='Intense Quote')
            doc.add_paragraph()
        
        # Add skills
        if resume_data.skills:
            doc.add_heading('Skills', level=2)
            skills_text = ', '.join(resume_data.skills)
            doc.add_paragraph(skills_text)
        
        # Save document
        doc.save(output_path)
        return output_path
